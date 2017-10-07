from docx import Document
import sys
import copy

# reload(sys)
# sys.setdefaultencoding('utf-8')

def change_paragraph(paragraph, filter_format):
    for format, content in filter_format.items():
        search_word = "{"+format+"}"
        needCheck = True if search_word in paragraph.text else False
        for run in paragraph.runs:
            if search_word in run.text:
                replaced = run.text.replace(search_word, content)
                run = run.clear()
                run.add_text(unicode(replaced))
                needCheck = False
        if needCheck:
            #print format, " is not checked in ", paragraph.text
            start_idx = 0
            end_idx = 0
            complete_word = ""
            i = 0
            broken_keys = []
            for run in paragraph.runs:
                if run.text == "{":
                    start_idx = i
                    complete_word = run.text
                else:
                    end_idx = i
                    complete_word = complete_word + run.text
                    if run.text == "}":
                        replaced = complete_word.replace(search_word, content)
                        broken_keys.append((start_idx, end_idx, replaced))
                #print ":::", run.text
                i = i+1

            for (start_idx, end_idx, replaced) in broken_keys:
                for run_idx, run in enumerate(paragraph.runs):
                    if start_idx <= run_idx and run_idx <= end_idx:
                        run = run.clear()
                    if run_idx == end_idx:
                        #print "))", replaced
                        run.add_text(replaced)


def replace(infile, outfile, filter_format):
    if filter_format == None:
        return

    document = Document(infile)

    for paragraph in document.paragraphs:
        change_paragraph(paragraph, filter_format)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    change_paragraph(paragraph, filter_format)

    document.save(outfile)

def main():
    replace("/Users/kchandra/Lyf/Kode/SCM/Github/k2/ResumeBuilder/demo/doc-template.docx", "output.docx", {"url1":"Alpa Beta"})

if __name__ == '__main__':
    main()